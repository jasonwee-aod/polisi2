from __future__ import annotations

from io import BytesIO
import pathlib
import sys

from docx import Document as DocxDocument
from openpyxl import Workbook

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1] / "src"))

from polisi_scraper.indexer.chunking import build_chunks
from polisi_scraper.indexer.parsers import DocxParser, HtmlParser, PdfParser, XlsxParser, get_parser


class _FakePdfPage:
    def __init__(self, text: str | None = None, *, raise_error: bool = False) -> None:
        self._text = text
        self._raise_error = raise_error

    def extract_text(self) -> str:
        if self._raise_error:
            raise ValueError("cannot extract")
        return self._text or ""


class _FakePdfReader:
    def __init__(self, *args: object, **kwargs: object) -> None:
        self.pages = [
            _FakePdfPage("Budget 2026 overview"),
            _FakePdfPage(raise_error=True),
            _FakePdfPage("Page 3 recovery text"),
        ]


def test_html_and_pdf_parsers_preserve_locators(monkeypatch) -> None:
    html_parser = HtmlParser()
    html = html_parser.parse_bytes(
        b"""
        <html>
          <head><title>Budget Highlights</title></head>
          <body>
            <h1>Subsidies</h1>
            <p>Fuel support continues in 2026.</p>
            <ul><li>Targeted households only.</li></ul>
          </body>
        </html>
        """
    )

    monkeypatch.setattr("polisi_scraper.indexer.parsers.pdf.PdfReader", _FakePdfReader)
    pdf = PdfParser().parse_bytes(b"%PDF-1.4", metadata={"title": "Budget PDF"})

    assert html.title == "Budget Highlights"
    assert html.blocks[0].section_heading == "Subsidies"
    assert html.blocks[1].block_type == "list_item"
    assert [block.page_number for block in pdf.blocks] == [1, 3]
    assert pdf.blocks[1].text == "Page 3 recovery text"


def test_docx_xlsx_and_chunking_preserve_structure() -> None:
    docx_buffer = BytesIO()
    doc = DocxDocument()
    doc.add_heading("Education Grants", level=1)
    doc.add_paragraph("Applicants must be enrolled full time.")
    doc.add_paragraph("Prepare income proof.", style="List Bullet")
    doc.save(docx_buffer)

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Allocations"
    sheet.append(["Program", "Amount"])
    sheet.append(["Bantuan Awal", "RM500"])
    xlsx_buffer = BytesIO()
    workbook.save(xlsx_buffer)

    parsed_docx = DocxParser().parse_bytes(docx_buffer.getvalue(), metadata={"title": "Grant Guide"})
    parsed_xlsx = XlsxParser().parse_bytes(xlsx_buffer.getvalue())

    chunks = build_chunks(parsed_docx, target_chars=40)

    assert parsed_docx.blocks[0].section_heading == "Education Grants"
    assert parsed_docx.blocks[1].block_type == "list_item"
    assert parsed_xlsx.blocks[0].sheet_name == "Allocations"
    assert parsed_xlsx.blocks[0].row_label == "Bantuan Awal"
    assert len(chunks) >= 2
    assert chunks[0].metadata["locators"][0]["section_heading"] == "Education Grants"
    assert any("Prepare income proof." in chunk.text for chunk in chunks)


def test_all_supported_file_types_parse(monkeypatch) -> None:
    monkeypatch.setattr("polisi_scraper.indexer.parsers.pdf.PdfReader", _FakePdfReader)

    docx_buffer = BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("DOCX content")
    doc.save(docx_buffer)

    workbook = Workbook()
    workbook.active.append(["Name", "Value"])
    workbook.active.append(["Threshold", "10"])
    xlsx_buffer = BytesIO()
    workbook.save(xlsx_buffer)

    payloads = {
        "html": b"<html><body><p>HTML content</p></body></html>",
        "pdf": b"%PDF-1.4",
        "docx": docx_buffer.getvalue(),
        "xlsx": xlsx_buffer.getvalue(),
    }

    for file_type, payload in payloads.items():
        parsed = get_parser(file_type).parse_bytes(payload)
        assert not parsed.is_empty()
