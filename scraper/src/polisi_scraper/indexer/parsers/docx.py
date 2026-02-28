"""DOCX parser preserving headings and list structure."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from polisi_scraper.indexer.parsers.base import DocumentParser, ParsedBlock, ParsedDocument


class DocxParser(DocumentParser):
    file_type = "docx"

    def parse_bytes(
        self,
        payload: bytes,
        *,
        metadata: dict[str, object] | None = None,
    ) -> ParsedDocument:
        document = Document(BytesIO(payload))
        blocks: list[ParsedBlock] = []
        current_heading: str | None = None

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower()
            if style_name.startswith("heading"):
                current_heading = text
                continue
            block_type = "list_item" if "list" in style_name else "paragraph"
            blocks.append(
                ParsedBlock(
                    text=text,
                    block_type=block_type,
                    section_heading=current_heading,
                )
            )

        return ParsedDocument(
            file_type=self.file_type,
            title=(metadata or {}).get("title") if metadata else None,
            blocks=blocks,
            metadata=dict(metadata or {}),
        )
